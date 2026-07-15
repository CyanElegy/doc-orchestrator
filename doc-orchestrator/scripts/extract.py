#!/usr/bin/env python3
"""
extract.py — Extract skeleton structure from a document template.

Usage:
    python3 extract.py <template> --output skeleton.json

Supported formats: .docx, .doc (via LibreOffice conversion), .xlsx, .xls

Output skeleton.json contains:
  - meta: document type, source file, warnings
  - units: ordered list of structural units (headings, paragraphs, tables, placeholders)
    Each unit has: id, type (fixed|placeholder|generated), element, para_index,
    and element-specific fields

Requires: python-docx, openpyxl, Pillow, pyyaml
"""

from __future__ import annotations

import json
import re
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path
from typing import Any, Optional

# ---------------------------------------------------------------------------
# Dependency checks
# ---------------------------------------------------------------------------

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("python-docx is not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(1)

try:
    import openpyxl
except ImportError:
    print("openpyxl is not installed. Run: pip install openpyxl", file=sys.stderr)
    sys.exit(1)

# ---------------------------------------------------------------------------
# Shared module import
# ---------------------------------------------------------------------------

try:
    from numbering_patterns import detect_numbering_prefix
except ImportError:
    from scripts.numbering_patterns import detect_numbering_prefix


# ---------------------------------------------------------------------------
# Patterns for placeholder detection
# ---------------------------------------------------------------------------

# 《...》 — explicit placeholder markers
RE_BOOK_PLACEHOLDER = re.compile(r"《([^》]+)》")

# 【...】 — template guidance text (should be replaced with real content)
RE_GUIDANCE_BRACKET = re.compile(r"【[^】]*】")

# XX...XX / ××...×× — implicit placeholder patterns
RE_X_PLACEHOLDER = re.compile(r"[X×]{2,}")

# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def extract(path: Path, output_path: Optional[Path] = None) -> dict[str, Any]:
    """Extract skeleton from a template file.

    Returns a dict (the skeleton). If output_path is given, writes JSON.
    """
    if not path.exists():
        return _error(f"File not found: {path}")

    suffix = path.suffix.lower()
    if suffix == ".doc":
        path = _convert_doc_to_docx(path)
        if path is None:
            return _error(f"Cannot process .doc file. LibreOffice is required for conversion.")
        suffix = ".docx"

    if suffix == ".docx":
        result = _extract_docx(path)
    elif suffix in (".xlsx", ".xls"):
        result = _extract_xlsx(path)
    else:
        return _error(f"Unsupported format: {suffix}. Supported: .docx, .doc, .xlsx, .xls")

    if output_path:
        output_path.write_text(
            json.dumps(result, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

    return result


# ---------------------------------------------------------------------------
# DOCX extraction
# ---------------------------------------------------------------------------

def _extract_docx(path: Path) -> dict[str, Any]:
    warnings: list[str] = []

    # Validate file
    try:
        with zipfile.ZipFile(path) as zf:
            if "[Content_Types].xml" not in zf.namelist():
                return _error(f"'{path}' is not a valid .docx file")
    except zipfile.BadZipFile:
        return _error(f"'{path}' is not a valid .docx file (not a ZIP archive)")

    try:
        doc = Document(str(path))
    except Exception as exc:
        return _error(f"Cannot open '{path}': {exc}")

    units: list[dict[str, Any]] = []
    unit_counter = 0

    def _next_id(prefix: str) -> str:
        nonlocal unit_counter
        unit_counter += 1
        return f"{prefix}-{unit_counter:03d}"

    # Check for SmartArt / text boxes (warn only)
    body_xml = doc.element.body.xml
    if "<w:txbxContent" in body_xml or "<mc:AlternateContent" in body_xml:
        warnings.append(
            "Template contains text boxes or SmartArt graphics. "
            "Text inside these elements cannot be automatically extracted. "
            "Please verify the skeleton output for missing content."
        )

    # Build body-index → paragraph-index mapping
    # (body children are interleaved <w:p> and <w:tbl> elements)
    body = doc.element.body
    body_idx_to_para_idx: dict[int, int] = {}
    body_idx_to_table_idx: dict[int, int] = {}
    para_count = 0
    table_count = 0

    for body_idx, child in enumerate(body):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            body_idx_to_para_idx[body_idx] = para_count
            para_count += 1
        elif tag == 'tbl':
            body_idx_to_table_idx[body_idx] = table_count
            table_count += 1

    # Collect headings first — we need to find the cover page boundary
    heading_paras: list[tuple[int, str, str, int, int]] = []  # (level, text, numbering, para_idx, body_idx)

    for idx, para in enumerate(doc.paragraphs):
        style = para.style
        if style and style.name and style.name.startswith("Heading "):
            try:
                level = int(style.name.split()[-1])
            except ValueError:
                continue
            raw_text = para.text.strip()
            if not raw_text:
                continue
            prefix, _ = detect_numbering_prefix(raw_text)
            clean_text = raw_text
            if prefix:
                clean_text = raw_text[len(prefix):].strip()
            # Find the body index for this paragraph
            body_idx = _find_body_index_for_para(body, para, body_idx_to_para_idx)
            heading_paras.append((level, clean_text, prefix or "", idx, body_idx))

    # Find first H1 — everything before it is the cover page
    first_h1_body_idx: Optional[int] = None
    for level, text, numbering, para_idx, body_idx in heading_paras:
        if level == 1:
            first_h1_body_idx = body_idx
            break

    # Process all body children in document order
    prev_text_context = ""  # track the last meaningful paragraph text for image context

    for body_idx, child in enumerate(body):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if tag == 'sectPr':
            # Section properties — skip (not a content unit)
            continue

        if tag == 'tbl':
            # Table element
            table_idx = body_idx_to_table_idx.get(body_idx, -1)
            if table_idx >= 0 and table_idx < len(doc.tables):
                table = doc.tables[table_idx]
                _extract_table_unit(table, table_idx, body_idx, units, _next_id)
            continue

        if tag != 'p':
            continue

        # --- Paragraph processing ---
        para_idx = body_idx_to_para_idx.get(body_idx, -1)
        if para_idx < 0 or para_idx >= len(doc.paragraphs):
            continue

        para = doc.paragraphs[para_idx]
        text = para.text.strip()
        style_name = para.style.name if para.style else ""

        in_cover = first_h1_body_idx is not None and body_idx < first_h1_body_idx

        # Handle empty paragraphs
        if not text:
            images_in_para = _find_images_in_paragraph(para)
            if images_in_para:
                # Look ahead for context after the image
                context_after = ""
                for ahead_idx in range(para_idx + 1, min(para_idx + 4, len(doc.paragraphs))):
                    ahead_text = doc.paragraphs[ahead_idx].text.strip()
                    ahead_style = doc.paragraphs[ahead_idx].style.name if doc.paragraphs[ahead_idx].style else ""
                    if ahead_text and not ahead_style.startswith("Heading "):
                        context_after = ahead_text[:120]
                        break
                    elif ahead_text and ahead_style.startswith("Heading "):
                        break

                for img_alt in images_in_para:
                    img_unit: dict[str, Any] = {
                        "id": _next_id("img"),
                        "type": "generated",
                        "element": "image",
                        "alt_text": img_alt,
                        "context_before": prev_text_context[-200:] if prev_text_context else "",
                        "context_after": context_after,
                        "para_index": body_idx,
                    }
                    is_informational = _image_is_informational(img_alt, prev_text_context, context_after)
                    if is_informational:
                        img_unit["image_type"] = "informational"
                        img_unit["warning"] = (
                            "此图片可能包含需要更新的项目特定内容"
                            "（项目名称、模块名、功能点等），建议确认后处理"
                        )
                    else:
                        img_unit["image_type"] = "decorative"
                        img_unit["warning"] = None
                    units.append(img_unit)
            elif in_cover:
                # Cover page spacer paragraph — preserve
                units.append({
                    "id": _next_id("p"),
                    "type": "fixed",
                    "element": "paragraph",
                    "text": "",
                    "para_index": body_idx,
                    "spacer": True,
                    "cover": True,
                })
            else:
                # Body empty paragraph — skip
                pass
            continue

        # --- Heading paragraphs ---
        if style_name.startswith("Heading "):
            try:
                level = int(style_name.split()[-1])
            except ValueError:
                level = 1
            prefix, _ = detect_numbering_prefix(text)
            clean = text[len(prefix):].strip() if prefix else text
            units.append({
                "id": _next_id("h"),
                "type": "fixed",
                "element": "heading",
                "level": level,
                "text": clean,
                "numbering": prefix or "",
                "para_index": body_idx,
            })
            prev_text_context = text
            continue

        # --- Body paragraph classification ---
        inline_images = _find_images_in_paragraph(para)

        # Priority 1: 《...》 placeholders
        book_placeholders = RE_BOOK_PLACEHOLDER.findall(text)
        if book_placeholders:
            for ph in book_placeholders:
                units.append({
                    "id": _next_id("ph"),
                    "type": "placeholder",
                    "element": "placeholder",
                    "pattern": f"《{ph}》",
                    "key": ph,
                    "para_index": body_idx,
                })
            remaining = RE_BOOK_PLACEHOLDER.sub("", text).strip()
            if remaining:
                units.append({
                    "id": _next_id("p"),
                    "type": "fixed",
                    "element": "paragraph",
                    "text": remaining,
                    "para_index": body_idx,
                })
        # Priority 2: 【...】 guidance text → always generated
        elif RE_GUIDANCE_BRACKET.search(text):
            # Guidance brackets mean this paragraph is a template instruction
            # Extract the guidance as description for the agent
            units.append({
                "id": _next_id("p"),
                "type": "generated",
                "element": "paragraph",
                "description": text,
                "para_index": body_idx,
            })
        # Priority 3: XX/×××× placeholder patterns
        elif RE_X_PLACEHOLDER.search(text):
            units.append({
                "id": _next_id("p"),
                "type": "generated",
                "element": "paragraph",
                "description": text,
                "para_index": body_idx,
                "has_x_placeholder": True,
            })
        # Priority 4: Cover page paragraphs → generated
        elif in_cover:
            units.append({
                "id": _next_id("p"),
                "type": "generated",
                "element": "paragraph",
                "description": text,
                "para_index": body_idx,
                "cover": True,
            })
        # Priority 5: Default — short text is generated, long text is fixed
        elif len(text) > 80:
            units.append({
                "id": _next_id("p"),
                "type": "fixed",
                "element": "paragraph",
                "text": text,
                "para_index": body_idx,
            })
        else:
            units.append({
                "id": _next_id("p"),
                "type": "generated",
                "element": "paragraph",
                "description": text if text else "正文内容",
                "para_index": body_idx,
            })

        # Add inline images found in this paragraph
        for img_alt in inline_images:
            img_unit: dict[str, Any] = {
                "id": _next_id("img"),
                "type": "generated",
                "element": "image",
                "alt_text": img_alt,
                "context_before": text[:200],
                "context_after": "",
                "inline": True,
                "para_index": body_idx,
            }
            if _image_is_informational(img_alt, text, ""):
                img_unit["image_type"] = "informational"
                img_unit["warning"] = (
                    "此图片可能包含需要更新的项目特定内容"
                    "（项目名称、模块名、功能点等），建议确认后处理"
                )
            else:
                img_unit["image_type"] = "decorative"
                img_unit["warning"] = None
            units.append(img_unit)

        prev_text_context = text

    return {
        "meta": {
            "type": "docx",
            "source": str(path.name),
            "extracted_at": None,
        },
        "warnings": warnings or None,
        "units": units,
        "unit_count": len(units),
        "summary": _build_summary(units),
    }


def _find_body_index_for_para(body, para, body_idx_to_para_idx: dict[int, int]) -> int:
    """Find the body child index for a given Paragraph object."""
    para_element = para._element
    for body_idx, child in enumerate(body):
        if child is para_element:
            return body_idx
    # Fallback: search through the mapping
    return -1


def _extract_table_unit(table, table_idx: int, body_idx: int,
                        units: list[dict[str, Any]], _next_id) -> None:
    """Extract a table unit with full cell content."""
    headers: list[str] = []
    data_rows: list[list[str]] = []

    for row_idx, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        if row_idx == 0:
            headers = cells
        else:
            data_rows.append(cells)

    units.append({
        "id": _next_id("t"),
        "type": "generated",
        "element": "table",
        "table_index": table_idx,
        "headers": headers,
        "rows": len(data_rows),
        "data": data_rows if data_rows else [],
        "para_index": body_idx,
    })


def _find_images_in_paragraph(para) -> list[str]:
    """Extract alt text from images in a paragraph. Returns list of alt text strings."""
    alts: list[str] = []
    namespace = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
    }
    for drawing in para._element.iter(qn("w:drawing")):
        c_nv_pr = drawing.find(".//" + qn("pic:cNvPr"), namespace)
        if c_nv_pr is not None:
            alt = c_nv_pr.get("descr", c_nv_pr.get("name", ""))
            if alt:
                alts.append(alt)
    return alts


def _image_is_informational(alt_text: str, before: str, after: str) -> bool:
    """Classify image as informational vs decorative using alt text and surrounding context."""
    combined = f"{alt_text} {before} {after}".lower()

    decorative_keywords = [
        "logo", "图标", "icon", "背景", "background",
        "分隔线", "divider", "装饰", "decoration",
    ]
    for kw in decorative_keywords:
        if kw in combined:
            return False

    informational_keywords = [
        "架构", "architecture", "拓扑", "topology",
        "流程", "flow", "flowchart", "部署", "deployment",
        "网络", "network", "数据流", "data flow",
        "组件", "component", "模块", "module",
        "结构", "structure", "示意图", "diagram",
        "框架", "framework", "层级", "layer", "tier",
        "系统", "system",
    ]
    for kw in informational_keywords:
        if kw in combined:
            return True

    return True


def _build_summary(units: list[dict]) -> dict[str, Any]:
    """Build a human-readable summary of the skeleton."""
    headings = [u for u in units if u["element"] == "heading"]
    tables = [u for u in units if u["element"] == "table"]
    placeholders = [u for u in units if u["element"] == "placeholder"]
    fixed = [u for u in units if u["type"] == "fixed"]
    generated = [u for u in units if u["type"] == "generated"]

    return {
        "total_units": len(units),
        "headings": len(headings),
        "tables": len(tables),
        "placeholders": len(placeholders),
        "fixed_units": len(fixed),
        "generated_units": len(generated),
    }


# ---------------------------------------------------------------------------
# XLSX extraction
# ---------------------------------------------------------------------------

def _extract_xlsx(path: Path) -> dict[str, Any]:
    warnings: list[str] = []

    try:
        if path.suffix.lower() == ".xls":
            wb = openpyxl.load_workbook(str(path), read_only=True)
        else:
            wb = openpyxl.load_workbook(str(path), read_only=True)
    except Exception as exc:
        return _error(f"Cannot open '{path}': {exc}")

    units: list[dict[str, Any]] = []
    unit_counter = 0

    def _next_id(prefix: str) -> str:
        nonlocal unit_counter
        unit_counter += 1
        return f"{prefix}-{unit_counter:03d}"

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]

        header_row_idx = _find_xlsx_header(ws)
        if header_row_idx is None:
            warnings.append(f"Sheet '{sheet_name}': could not detect header row, skipping")
            continue

        headers: list[str] = []
        for col in range(1, ws.max_column + 1):
            val = ws.cell(row=header_row_idx, column=col).value
            headers.append(str(val).strip() if val is not None else f"Column{col}")

        data_rows = 0
        for row_idx in range(header_row_idx + 1, ws.max_row + 1):
            if any(
                ws.cell(row=row_idx, column=col).value is not None
                for col in range(1, ws.max_column + 1)
            ):
                data_rows += 1

        units.append({
            "id": _next_id("sheet"),
            "type": "generated" if data_rows == 0 else "fixed",
            "element": "sheet",
            "sheet_name": sheet_name,
            "header_row": header_row_idx,
            "headers": headers,
            "data_rows": data_rows,
        })

    wb.close()

    return {
        "meta": {
            "type": "xlsx",
            "source": str(path.name),
        },
        "warnings": warnings or None,
        "units": units,
        "unit_count": len(units),
    }


def _find_xlsx_header(ws, max_check: int = 20) -> Optional[int]:
    """Find the header row in an Excel sheet."""
    max_col = ws.max_column
    for row_idx in range(1, min(max_check, ws.max_row) + 1):
        non_empty = 0
        for col_idx in range(1, max_col + 1):
            if ws.cell(row=row_idx, column=col_idx).value is not None:
                non_empty += 1
        if non_empty >= max(1, max_col // 2):
            return row_idx
    return 1 if ws.max_row > 0 else None


# ---------------------------------------------------------------------------
# .doc → .docx conversion
# ---------------------------------------------------------------------------

def _convert_doc_to_docx(path: Path) -> Optional[Path]:
    """Convert a .doc file to .docx using LibreOffice. Returns the .docx path or None."""
    try:
        result = subprocess.run(
            [
                "soffice",
                "--headless",
                "--convert-to", "docx",
                "--outdir", str(path.parent),
                str(path),
            ],
            capture_output=True,
            timeout=60,
        )
        if result.returncode != 0:
            print(
                f"Warning: LibreOffice conversion failed: {result.stderr.decode()}",
                file=sys.stderr,
            )
            return None
        docx_path = path.with_suffix(".docx")
        if docx_path.exists():
            return docx_path
        return None
    except FileNotFoundError:
        print(
            "Warning: LibreOffice not found. Cannot convert .doc files. "
            "Install LibreOffice: brew install --cask libreoffice",
            file=sys.stderr,
        )
        return None
    except Exception as exc:
        print(f"Warning: .doc conversion failed: {exc}", file=sys.stderr)
        return None


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _error(message: str) -> dict[str, Any]:
    """Return a structured error response."""
    return {
        "meta": {"type": "error"},
        "error": True,
        "message": message,
        "units": [],
        "unit_count": 0,
    }


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main() -> None:
    args = sys.argv[1:]
    if not args or args[0] in ("-h", "--help"):
        print(__doc__)
        sys.exit(0)

    template_path = Path(args[0])
    output_path: Optional[Path] = None

    i = 1
    while i < len(args):
        if args[i] == "--output" and i + 1 < len(args):
            output_path = Path(args[i + 1])
            i += 2
        else:
            i += 1

    result = extract(template_path, output_path)

    if result.get("error"):
        print(f"Error: {result['message']}", file=sys.stderr)
        sys.exit(1)

    if not output_path:
        print(json.dumps(result, ensure_ascii=False, indent=2))

    summary = result.get("summary", {})
    warning_count = len(result.get("warnings") or [])
    status = " ⚠" if warning_count > 0 else ""
    print(
        f"\n✓ {result['unit_count']} units extracted{status}",
        file=sys.stderr,
    )
    if warning_count:
        for w in result["warnings"]:
            print(f"  ⚠ {w}", file=sys.stderr)


if __name__ == "__main__":
    main()
