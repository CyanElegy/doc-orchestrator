#!/usr/bin/env python3
"""
assemble_docx.py — Assemble a styled .docx from AI-generated Markdown and a template.

Usage:
    python3 assemble_docx.py <content.md> <template.docx> <project.yaml> <output.docx>

Steps:
    1. Copy template.docx to output.docx (inherit styles, headers, footers, page setup).
    2. Read content.md — parse YAML front matter + Markdown body.
    3. Read project.yaml — key-value pairs for ${key} substitution.
    4. Clear body content from output.docx (keep sectPr).
    5. Parse Markdown into Word sections with numbering from template.
    6. Replace ${key} placeholders in headers/footers with project.yaml values.
    7. Insert TOC field code at document start.
    8. Save.

Requires: python-docx, Pillow (PIL), pyyaml.
"""

from __future__ import annotations

import re
import shutil
import sys
from pathlib import Path
from typing import Any, Optional

# ---------------------------------------------------------------------------
# Dependency checks — fail early with clear messages
# ---------------------------------------------------------------------------

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Emu, RGBColor
    from docx.text.paragraph import Paragraph
    from docx.table import Table as DocxTable
except ImportError:
    print(
        "python-docx is not installed. Run: pip install python-docx",
        file=sys.stderr,
    )
    sys.exit(1)

try:
    import yaml
except ImportError:
    print(
        "pyyaml is not installed. Run: pip install pyyaml",
        file=sys.stderr,
    )
    sys.exit(1)

try:
    from PIL import Image as PILImage
except ImportError:
    print(
        "Pillow (PIL) is not installed. Run: pip install Pillow",
        file=sys.stderr,
    )
    sys.exit(1)

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

HEADING_STYLE_MAP: dict[int, str] = {
    1: "Heading 1",
    2: "Heading 2",
    3: "Heading 3",
    4: "Heading 4",
    5: "Heading 5",
    6: "Heading 6",
    7: "Heading 7",
    8: "Heading 8",
    9: "Heading 9",
}

# Numbering regex patterns: (compiled_pattern, label)
_NUMBERING_PATTERNS: list[tuple[re.Pattern[str], str]] = [
    (re.compile(r"^[一二三四五六七八九十百千]+[、．]?\s*"), "一、"),
    (re.compile(r"^（[一二三四五六七八九十百千]+）\s*"), "（一）"),
    (re.compile(r"^\d+\.[\d.]*\s*"), "1.1"),
    (re.compile(r"^第[一二三四五六七八九十百千]+[章节条]\s*"), "第一章"),
    (re.compile(r"^\(\d+\)\s*"), "(1)"),
    (re.compile(r"^\d+[、．.]\s*"), "1、"),
    (re.compile(r"^[①②③④⑤⑥⑦⑧⑨⑩]\s*"), "①"),
]


def _detect_numbering_prefix(text: str) -> str:
    """Detect numbering prefix in heading text.

    Returns the matched prefix string (e.g. '一、', '（一）', '1.') or ''.
    """
    for pattern, _label in _NUMBERING_PATTERNS:
        m = pattern.match(text)
        if m:
            return m.group().strip()
    return ""


# ---------------------------------------------------------------------------
# Markdown parsing
# ---------------------------------------------------------------------------


def parse_markdown(filepath: Path) -> tuple[dict[str, Any], str]:
    """Parse a Markdown file with optional YAML front matter.

    Returns (front_matter_dict, body_string).
    Raises ValueError on malformed YAML.
    """
    content = filepath.read_text(encoding="utf-8")

    front_matter: dict[str, Any] = {}
    body = content

    if content.startswith("---"):
        end_idx = content.find("---", 3)
        if end_idx == -1:
            raise ValueError(
                f"{filepath}: YAML front matter delimiter '---' not closed"
            )
        yaml_block = content[3:end_idx].strip()
        body = content[end_idx + 3:].strip()

        if yaml_block:
            try:
                parsed = yaml.safe_load(yaml_block)
                if isinstance(parsed, dict):
                    front_matter = parsed
                elif parsed is not None:
                    print(
                        f"Warning: {filepath}: YAML front matter is not a dict "
                        f"(type: {type(parsed).__name__}), treating as empty",
                        file=sys.stderr,
                    )
            except yaml.YAMLError as exc:
                line_num = ""
                if hasattr(exc, "problem_mark") and exc.problem_mark is not None:
                    line_num = f" (near line {exc.problem_mark.line + 1})"
                raise ValueError(
                    f"{filepath}: Malformed YAML{line_num}: {exc}"
                ) from exc

    return front_matter, body


def _split_markdown_blocks(body: str) -> list[dict[str, Any]]:
    """Split Markdown body into a list of block dicts.

    Each block has a 'type' key: 'heading', 'paragraph', 'table',
    'thematic_break'.
    """
    lines = body.split("\n")
    blocks: list[dict[str, Any]] = []
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Thematic break
        if stripped in ("---", "***", "___"):
            blocks.append({"type": "thematic_break"})
            i += 1
            continue

        # Heading
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            blocks.append({"type": "heading", "level": min(level, 9), "text": text})
            i += 1
            continue

        # Table — collect consecutive table rows
        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines: list[str] = []
            table_lines.append(stripped)
            i += 1
            # Skip separator row if present
            if i < len(lines) and re.match(r"^[\s|:\-]+$", lines[i].strip()):
                table_lines.append(lines[i].strip())
                i += 1
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            blocks.append({"type": "table", "lines": table_lines})
            continue

        # Empty line
        if not stripped:
            i += 1
            continue

        # Regular paragraph (multi-line)
        para_lines: list[str] = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt:
                break
            if re.match(r"^#{1,6}\s+", nxt):
                break
            if nxt.startswith("|") and nxt.endswith("|"):
                break
            if nxt in ("---", "***", "___"):
                break
            para_lines.append(nxt)
            i += 1

        para_text = " ".join(para_lines)
        if para_text:
            blocks.append({"type": "paragraph", "text": para_text})

    return blocks


# ---------------------------------------------------------------------------
# Numbering pattern detection (from template)
# ---------------------------------------------------------------------------

NumberingPattern = tuple[int, str, str]  # (level, prefix_example, style_label)


def detect_template_numbering(template_doc: Document) -> list[NumberingPattern]:
    """Scan template for heading numbering patterns.

    Returns [(level, prefix_example, style_label), ...],
    e.g. [(1, '一、', 'Heading 1'), (2, '（一）', 'Heading 2')].
    """
    patterns: list[NumberingPattern] = []
    processed_levels: set[int] = set()

    for para in template_doc.paragraphs:
        style = para.style
        if not style or not style.name or not style.name.startswith("Heading "):
            continue
        try:
            level = int(style.name.split()[-1])
        except ValueError:
            continue
        if level in processed_levels:
            continue

        text = para.text.strip()
        if not text:
            continue

        prefix = _detect_numbering_prefix(text)
        if prefix:
            patterns.append((level, prefix, style.name))
            processed_levels.add(level)

    return sorted(patterns, key=lambda x: x[0])


# ---------------------------------------------------------------------------
# Document assembly helpers
# ---------------------------------------------------------------------------


def _copy_docx(src: Path, dst: Path) -> None:
    """Copy a .docx file preserving all internal parts."""
    shutil.copy2(str(src), str(dst))


def _clear_body(doc: Document) -> None:
    """Remove all paragraphs and tables from the document body, keeping sectPr."""
    body = doc.element.body
    sect_pr = body.find(qn("w:sectPr"))

    children = list(body)
    for child in children:
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)

    if sect_pr is not None:
        body.append(sect_pr)


def _get_style_or_fallback(doc: Document, level: int) -> Optional[str]:
    """Get heading style name for *level*, falling back to next available."""
    for lvl in range(level, 0, -1):
        sname = HEADING_STYLE_MAP.get(lvl, "")
        try:
            doc.styles[sname]
            return sname
        except KeyError:
            continue
    return None


def _add_formatted_paragraph(
    doc: Document,
    text: str,
    style_name: str = "Normal",
) -> Paragraph:
    """Add a paragraph with the given style."""
    para = doc.add_paragraph()
    para.style = doc.styles[style_name]
    para.add_run(text)
    return para


def _add_heading_with_numbering(
    doc: Document,
    text: str,
    level: int,
    numbering: dict[int, str],
) -> Paragraph:
    """Add a heading paragraph, optionally prepending a numbering prefix."""
    style_name = _get_style_or_fallback(doc, level)
    if style_name is None:
        style_name = "Normal"

    prefix = numbering.get(level, "")
    full_text = f"{prefix}{text}" if prefix else text

    para = doc.add_paragraph()
    para.style = doc.styles[style_name]
    para.add_run(full_text)
    return para


def _replace_placeholders(text: str, project_vars: dict[str, str]) -> str:
    """Replace ${key} with values from project_vars dict."""
    def _replacer(m: re.Match[str]) -> str:
        return project_vars.get(m.group(1), m.group(0))
    return re.sub(r"\$\{([^}]+)\}", _replacer, text)


def _substitute_placeholders_in_para(
    para: Paragraph, project_vars: dict[str, str]
) -> None:
    """Replace ${key} patterns in all runs of a paragraph."""
    for run in para.runs:
        if run.text:
            run.text = _replace_placeholders(run.text, project_vars)


def _embed_image(para: Paragraph, filename: str, doc: Document) -> None:
    """Embed an image file into a paragraph, auto-sizing to page width."""
    candidates = [Path(filename)]
    if not candidates[0].is_absolute():
        candidates.append(Path.cwd() / filename)

    img_path: Optional[Path] = None
    for cand in candidates:
        if cand.exists():
            img_path = cand
            break

    if img_path is None:
        print(
            f"Warning: Image not found: {filename}. Inserting placeholder.",
            file=sys.stderr,
        )
        run = para.add_run(f"[Image: {filename}]")
        run.font.italic = True
        try:
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        except Exception:
            pass
        return

    try:
        sect = doc.sections[0]
        page_width = sect.page_width or Emu(11907000)
        left_margin = sect.left_margin or Emu(1440000)
        right_margin = sect.right_margin or Emu(1440000)
        avail_emu = page_width - left_margin - right_margin
        avail_px = avail_emu / 914400 * 96

        with PILImage.open(img_path) as pil_img:
            img_w, img_h = pil_img.size

        if img_w > avail_px:
            ratio = avail_px / img_w
            display_w = int(img_w * ratio * 914400 / 96)
        else:
            display_w = int(img_w * 914400 / 96)

        run = para.add_run()
        run.add_picture(str(img_path), width=Emu(display_w))
    except Exception as exc:
        print(
            f"Warning: Could not embed image '{filename}': {exc}. "
            f"Inserting placeholder.",
            file=sys.stderr,
        )
        run = para.add_run(f"[Image: {filename}]")
        run.font.italic = True


def _process_inline_markdown(
    para: Paragraph,
    text: str,
    doc: Document,
    project_vars: dict[str, str],
) -> None:
    """Process inline {{...}} macros in paragraph text.

    Supported: {{asset:path}} embeds an image; {{ref:target}} inserts a
    cross-reference text.  Also substitutes ${key} placeholders.
    """
    para.clear()

    pattern = re.compile(r"\{\{(asset|ref):([^}]+)\}\}")
    last_end = 0
    for m in pattern.finditer(text):
        before = text[last_end:m.start()]
        if before:
            para.add_run(before)

        macro_type = m.group(1)
        macro_arg = m.group(2).strip()

        if macro_type == "asset":
            _embed_image(para, macro_arg, doc)
        elif macro_type == "ref":
            run = para.add_run(f"[{macro_arg}]")
            try:
                run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
            except Exception:
                pass
            run.font.underline = True

        last_end = m.end()

    trailing = text[last_end:]
    if trailing:
        para.add_run(trailing)

    _substitute_placeholders_in_para(para, project_vars)


def _create_table_from_lines(doc: Document, table_lines: list[str]) -> DocxTable:
    """Create a Word table from Markdown-style pipe table lines."""
    if len(table_lines) == 0:
        return doc.add_table(rows=1, cols=1)

    # First line is header; second (if pipe-only) is separator; rest are data rows
    header = [c.strip() for c in table_lines[0].strip().strip("|").split("|")]
    rows_data: list[list[str]] = [header]
    start = 2 if len(table_lines) > 1 and re.match(r"^[\s|:\-]+$", table_lines[1]) else 1
    for line in table_lines[start:]:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        while len(cells) < len(header):
            cells.append("")
        rows_data.append(cells)

    num_cols = max(len(r) for r in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    try:
        table.style = doc.styles["Table Grid"]
    except KeyError:
        pass

    for row_idx, row_cells in enumerate(rows_data):
        for col_idx in range(num_cols):
            cell_text = row_cells[col_idx] if col_idx < len(row_cells) else ""
            table.cell(row_idx, col_idx).text = cell_text

    return table


def _insert_toc(doc: Document) -> None:
    """Insert TOC heading and TOC field code at the very start of the document."""
    # Because the body is empty after _clear_body, add paragraphs in reverse
    # order-of-appearance then move them to front.

    # 1) TOC heading paragraph
    toc_heading_para = doc.add_paragraph()
    toc_heading_para.style = doc.styles["Heading 1"]
    run = toc_heading_para.add_run("目  录")

    # 2) TOC field paragraph (empty, will hold field codes)
    toc_field_para = doc.add_paragraph()

    # Build TOC field XML
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.set(qn("xml:space"), "preserve")
    placeholder.text = '[请在 Word 中更新目录 — 右键点击此处选择“更新域”]'

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r1 = toc_field_para.add_run()
    r1._element.append(fld_begin)
    r2 = toc_field_para.add_run()
    r2._element.append(instr)
    r3 = toc_field_para.add_run()
    r3._element.append(fld_sep)
    r4 = toc_field_para.add_run()
    r4._element.append(placeholder)
    r5 = toc_field_para.add_run()
    r5._element.append(fld_end)

    # 3) Page-break paragraph after TOC
    pb_para = doc.add_paragraph()
    pb_run = pb_para.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    pb_run._element.append(br)

    # Move the three paragraphs to the front of the body in order
    body = doc.element.body
    body.remove(toc_heading_para._element)
    body.remove(toc_field_para._element)
    body.remove(pb_para._element)
    body.insert(0, pb_para._element)
    body.insert(0, toc_field_para._element)
    body.insert(0, toc_heading_para._element)


def _substitute_header_footer(
    doc: Document, project_vars: dict[str, str]
) -> None:
    """Replace ${key} placeholders in headers and footers."""
    for section in doc.sections:
        for hf in (section.header, section.footer):
            if hf and not hf.is_linked_to_previous:
                for para in hf.paragraphs:
                    _substitute_placeholders_in_para(para, project_vars)


# ---------------------------------------------------------------------------
# Main assembly
# ---------------------------------------------------------------------------


def assemble(
    content_path: Path,
    template_path: Path,
    project_yaml_path: Path,
    output_path: Path,
) -> None:
    """Main assembly workflow."""
    # -- Validate inputs -------------------------------------------------------
    for p, label in [
        (content_path, "Content file"),
        (template_path, "Template file"),
        (project_yaml_path, "Project YAML"),
    ]:
        if not p.exists():
            print(f"Error: {label} not found: {p}", file=sys.stderr)
            sys.exit(1)

    if template_path.suffix.lower() != ".docx":
        print(
            f"Error: Template file must be .docx: {template_path}",
            file=sys.stderr,
        )
        sys.exit(1)

    if output_path.suffix.lower() != ".docx":
        print(
            f"Error: Output file must be .docx: {output_path}",
            file=sys.stderr,
        )
        sys.exit(1)

    # -- Step 1: Parse inputs --------------------------------------------------
    print(f"Reading content: {content_path}")
    try:
        _front_matter, body = parse_markdown(content_path)
    except ValueError as exc:
        print(f"Error: {exc}", file=sys.stderr)
        sys.exit(1)
    except Exception as exc:
        print(
            f"Error reading content file '{content_path}': {exc}",
            file=sys.stderr,
        )
        sys.exit(1)

    print(f"Reading project YAML: {project_yaml_path}")
    try:
        with open(project_yaml_path, encoding="utf-8") as f:
            project_data = yaml.safe_load(f)
    except yaml.YAMLError as exc:
        line_num = ""
        if hasattr(exc, "problem_mark") and exc.problem_mark is not None:
            line_num = f" (near line {exc.problem_mark.line + 1})"
        print(
            f"Error: Malformed YAML in '{project_yaml_path}'{line_num}: {exc}",
            file=sys.stderr,
        )
        sys.exit(1)
    except Exception as exc:
        print(
            f"Error reading '{project_yaml_path}': {exc}",
            file=sys.stderr,
        )
        sys.exit(1)

    if not isinstance(project_data, dict):
        print(
            f"Warning: '{project_yaml_path}' does not contain a dict, "
            f"treating as empty",
            file=sys.stderr,
        )
        project_vars: dict[str, str] = {}
    else:
        project_vars = {str(k): str(v) for k, v in project_data.items()}

    # -- Step 2: Copy template to output --------------------------------------
    print(f"Copying template to output: {output_path}")
    _copy_docx(template_path, output_path)

    # -- Step 3: Open output document and detect numbering --------------------
    doc = Document(str(output_path))

    template_doc = Document(str(template_path))
    numbering_patterns = detect_template_numbering(template_doc)
    numbering_map: dict[int, str] = {
        level: prefix for level, prefix, _label in numbering_patterns
    }
    # template_doc is no longer needed — Document objects don't require closing

    # -- Step 4: Clear body content -------------------------------------------
    print("Clearing body content (keeping sectPr)...")
    _clear_body(doc)

    # -- Step 5: Parse Markdown into blocks -----------------------------------
    blocks = _split_markdown_blocks(body)
    print(f"Parsed {len(blocks)} blocks from Markdown body")

    # -- Step 6: Write content into document ----------------------------------
    table_count = 0
    heading_count = 0
    for block in blocks:
        try:
            if block["type"] == "heading":
                heading_count += 1
                _add_heading_with_numbering(
                    doc, block["text"], block["level"], numbering_map
                )
            elif block["type"] == "paragraph":
                para = _add_formatted_paragraph(doc, block["text"])
                _process_inline_markdown(
                    para, block["text"], doc, project_vars
                )
            elif block["type"] == "table":
                table_count += 1
                _create_table_from_lines(doc, block["lines"])
            elif block["type"] == "thematic_break":
                para = doc.add_paragraph()
                run = para.add_run("_" * 40)
                try:
                    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
                except Exception:
                    pass
        except Exception as exc:
            print(
                f"Warning: Failed to process {block.get('type', 'unknown')} "
                f"block: {exc}",
                file=sys.stderr,
            )

    # -- Step 7: Insert TOC at the beginning ----------------------------------
    print("Inserting TOC field...")
    _insert_toc(doc)

    # -- Step 8: Substitute placeholders in headers/footers -------------------
    print("Substituting placeholders in headers/footers...")
    _substitute_header_footer(doc, project_vars)

    # -- Step 9: Save ---------------------------------------------------------
    print(f"Saving output: {output_path}")
    doc.save(str(output_path))

    print(f"Done! {heading_count} headings, {table_count} tables inserted.")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------


def main() -> None:
    """CLI entry point."""
    args = sys.argv[1:]
    if not args or args[0] in ("-h", "--help"):
        print(__doc__)
        sys.exit(0)

    if len(args) < 4:
        print(
            "Usage: python3 assemble_docx.py "
            "<content.md> <template.docx> <project.yaml> <output.docx>",
            file=sys.stderr,
        )
        sys.exit(1)

    assemble(
        Path(args[0]),
        Path(args[1]),
        Path(args[2]),
        Path(args[3]),
    )


if __name__ == "__main__":
    main()
