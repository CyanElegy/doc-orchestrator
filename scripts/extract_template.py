#!/usr/bin/env python3
"""
extract_template.py — Extract structural information from a .docx template file.

Usage:
    python3 extract_template.py <template.docx> [--output structure.json]

Extracts:
    1. Heading tree — nested headings with numbering prefix and guidance text.
    2. Style definitions — font, size, bold, alignment, indentation.
    3. Page setup — margins, page size.
    4. Header/footer text — default section header and footer.
    5. Placeholders — 《...》 patterns in body, headers, footers.
    6. Embedded images — alt-text and content type for inline images.

Requires: python-docx (pip install python-docx)
"""

from __future__ import annotations

import json
import re
import sys
import zipfile
from dataclasses import dataclass, field, asdict
from pathlib import Path
from typing import Any, Optional

try:
    from docx import Document
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print(
        "python-docx is not installed. Run: pip install python-docx",
        file=sys.stderr,
    )
    sys.exit(1)

# ---------------------------------------------------------------------------
# Data models
# ---------------------------------------------------------------------------


@dataclass
class NumberingInfo:
    prefix: str = ""
    pattern: str = ""


@dataclass
class HeadingNode:
    level: int
    text: str
    numbering_prefix: str = ""
    guidance: list[str] = field(default_factory=list)
    children: list[HeadingNode] = field(default_factory=list)


@dataclass
class StyleInfo:
    style_id: str
    font_name: str = ""
    font_size: Optional[int] = None  # half-points
    bold: Optional[bool] = None
    alignment: str = "left"
    first_line_indent: int = 0  # twips (converted from EMU)
    before_spacing: int = 0  # EMU
    after_spacing: int = 0  # EMU


@dataclass
class MarginInfo:
    top: int = 0
    bottom: int = 0
    left: int = 0
    right: int = 0


@dataclass
class PageSize:
    width: int = 0
    height: int = 0


@dataclass
class PageSetup:
    margins: MarginInfo = field(default_factory=MarginInfo)
    page_size: PageSize = field(default_factory=PageSize)


@dataclass
class HeaderFooter:
    header: str = ""
    footer: str = ""


@dataclass
class ImageInfo:
    alt_text: str = ""
    content_type: str = ""


@dataclass
class TemplateStructure:
    headings: list[HeadingNode] = field(default_factory=list)
    styles: dict[str, StyleInfo] = field(default_factory=dict)
    page_setup: PageSetup = field(default_factory=PageSetup)
    header_footer: HeaderFooter = field(default_factory=HeaderFooter)
    placeholders: list[str] = field(default_factory=list)
    images: list[ImageInfo] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Serialisation helpers
# ---------------------------------------------------------------------------


def _to_dict(obj: Any) -> Any:
    """Recursively convert dataclass tree to plain dicts."""
    if isinstance(obj, (list, tuple)):
        return [_to_dict(v) for v in obj]
    if hasattr(obj, "__dataclass_fields__"):
        return {k: _to_dict(v) for k, v in asdict(obj).items()}
    return obj


# ---------------------------------------------------------------------------
# Numbering prefix detection
# ---------------------------------------------------------------------------

# Recognised numbering patterns in order of priority
NUMBERING_PATTERNS: list[tuple[re.Pattern[str], str]] = [
    (re.compile(r"^[一二三四五六七八九十百千]+[、．]?\s*"), "一、"),  # 中文数字
    (re.compile(r"^（[一二三四五六七八九十百千]+）\s*"), "（一）"),  # （一）
    (re.compile(r"^\d+\.[\d.]*\s*"), "1.1"),  # 1.1 / 1.1.1
    (re.compile(r"^第[一二三四五六七八九十百千]+[章节条]\s*"), "第一章"),  # 第X章
    (re.compile(r"^\(\d+\)\s*"), "(1)"),  # (1)
    (re.compile(r"^\d+[、．.]\s*"), "1、"),  # 1、 / 1．
    (re.compile(r"^①\s*"), "①"),  # circled digit
    (re.compile(r"^[①②③④⑤⑥⑦⑧⑨⑩]\s*"), "①"),  # more circled digits
]


def detect_numbering_prefix(text: str) -> tuple[str, str]:
    """Detect numbering prefix in heading text.

    Returns (prefix, pattern_name) where pattern_name is a human-readable
    label like '一、' or '1.' indicating the numbering style.
    """
    for pattern, label in NUMBERING_PATTERNS:
        m = pattern.match(text)
        if m:
            return m.group().strip(), label
    return "", ""


def strip_numbering(text: str) -> str:
    """Remove numbering prefix from heading text."""
    for pattern, _label in NUMBERING_PATTERNS:
        m = pattern.match(text)
        if m:
            return text[m.end():].strip()
    return text.strip()


# ---------------------------------------------------------------------------
# Extraction helpers
# ---------------------------------------------------------------------------


def _get_or_none(obj: Any, attr: str) -> Any:
    """Safely get an attribute or None."""
    return getattr(obj, attr, None)


def _to_int(val: Any) -> int:
    """Convert a value to int; return 0 on failure."""
    if val is None:
        return 0
    try:
        return int(val)
    except (ValueError, TypeError):
        return 0


def _emu_to_twips(emu: Any) -> int:
    """Convert EMU to twips. 1 twip = 635 EMU."""
    return _to_int(emu) // 635


def _extract_text_from_element(el: Any) -> str:
    """Extract plain text from an oxml element."""
    texts: list[str] = []
    for t in el.iter(qn("w:t")):
        if t.text:
            texts.append(t.text)
    return "".join(texts)


def _extract_images_from_document(doc: Document) -> list[ImageInfo]:
    """Extract inline images from the document body."""
    images: list[ImageInfo] = []
    namespace = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
    }

    for drawing in doc.element.iter(qn("w:drawing")):
        alt_text = ""
        # Try wp:inline/wp:anchor -> a:graphic -> pic:pic -> pic:nvPicPr -> pic:cNvPr
        inline = drawing.find(".//" + qn("wp:inline"), namespace)
        if inline is None:
            inline = drawing.find(".//" + qn("wp:anchor"), namespace)
        if inline is not None:
            c_nv_pr = inline.find(".//" + qn("pic:cNvPr"), namespace)
            if c_nv_pr is not None:
                alt_text = c_nv_pr.get("descr", c_nv_pr.get("name", ""))

        # Content type from relationship
        blip = drawing.find(".//" + qn("a:blip"), namespace)
        content_type = ""
        if blip is not None:
            embed_id = blip.get(qn("r:embed"), "")
            if embed_id and doc.part.rels:
                rel = doc.part.rels.get(embed_id)
                if rel is not None:
                    content_type = rel.reltype or ""
                    # Shorten common URIs
                    if "jpeg" in content_type:
                        content_type = "image/jpeg"
                    elif "png" in content_type:
                        content_type = "image/png"
                    elif "gif" in content_type:
                        content_type = "image/gif"
                    elif "bmp" in content_type:
                        content_type = "image/bmp"
                    elif "wmf" in content_type:
                        content_type = "image/wmf"

        images.append(ImageInfo(alt_text=alt_text, content_type=content_type))

    return images


# ---------------------------------------------------------------------------
# Main extraction
# ---------------------------------------------------------------------------


def extract_template(path: Path) -> TemplateStructure:
    """Extract full structure from a .docx template file."""
    if not path.exists():
        print(f"Error: File not found: {path}", file=sys.stderr)
        sys.exit(1)

    if path.suffix.lower() not in (".docx",):
        print(
            f"Error: '{path}' is not a .docx file (suffix: {path.suffix})",
            file=sys.stderr,
        )
        sys.exit(1)

    # Quick check: is it really a DOCX (ZIP archive with [Content_Types].xml)?
    try:
        with zipfile.ZipFile(path) as zf:
            if "[Content_Types].xml" not in zf.namelist():
                print(
                    f"Error: '{path}' is not a valid .docx file (missing [Content_Types].xml)",
                    file=sys.stderr,
                )
                sys.exit(1)
    except zipfile.BadZipFile:
        print(
            f"Error: '{path}' is not a valid .docx file (not a ZIP archive)",
            file=sys.stderr,
        )
        sys.exit(1)

    doc = Document(str(path))

    struct = TemplateStructure()

    # -- Detect empty document --------------------------------------------------
    body_elements = list(doc.element.body)
    if len(body_elements) == 0 or (
        len(body_elements) == 1
        and body_elements[0].tag == qn("w:sectPr")
    ):
        print("Warning: Document appears to be empty.", file=sys.stderr)
        return struct

    # -- 1. Heading tree --------------------------------------------------------
    all_paragraphs = list(doc.paragraphs)
    heading_paras: list[tuple[int, str, str, int]] = []  # (level, text, numbering, index)

    for idx, para in enumerate(all_paragraphs):
        style = para.style
        if style and style.name and style.name.startswith("Heading "):
            try:
                level = int(style.name.split()[-1])
            except ValueError:
                continue
            raw_text = para.text.strip()
            if not raw_text:
                continue
            number_prefix, _pattern = detect_numbering_prefix(raw_text)
            clean_text = strip_numbering(raw_text)
            heading_paras.append((level, clean_text, number_prefix, idx))

    # Build nested tree
    root: list[HeadingNode] = []
    stack: list[tuple[int, HeadingNode]] = []  # (level, node)

    for level, clean_text, number_prefix, idx in heading_paras:
        node = HeadingNode(level=level, text=clean_text, numbering_prefix=number_prefix)

        # Pop stack until we find parent
        while stack and stack[-1][0] >= level:
            stack.pop()

        if stack:
            stack[-1][1].children.append(node)
        else:
            root.append(node)

        stack.append((level, node))

    # Attach guidance (body paragraphs between headings)
    for i, (level, clean_text, number_prefix, idx) in enumerate(heading_paras):
        # Find the node we just built (walk tree)
        target_node = _find_node(root, level, clean_text)
        if target_node is None:
            continue

        # Determine range: after this heading up to next heading of same or higher level
        next_boundary = len(all_paragraphs)
        for j, (nlvl, _ctxt, _np, nidx) in enumerate(heading_paras[i + 1:], start=i + 1):
            if nlvl <= level:
                next_boundary = nidx
                break

        start_idx = idx + 1
        if start_idx < next_boundary:
            guidance_paras: list[str] = []
            for gidx in range(start_idx, next_boundary):
                p = all_paragraphs[gidx]
                txt = p.text.strip()
                if txt and p.style.name != "Heading 1":
                    guidance_paras.append(txt)
            target_node.guidance = guidance_paras

    struct.headings = root

    # -- 2. Style definitions ---------------------------------------------------
    used_style_ids: set[str] = set()
    for para in doc.paragraphs:
        if para.style:
            used_style_ids.add(para.style.name)

    for style in doc.styles:
        if style.name not in used_style_ids:
            continue
        font = style.font
        pf = style.paragraph_format

        font_name = ""
        if font and font.name:
            font_name = font.name

        font_size: Optional[int] = None
        if font and font.size:
            # font.size is in EMU; 1 pt = 12700 EMU → half-points = EMU / 6350
            font_size = font.size // 6350

        bold: Optional[bool] = None
        if font:
            bold = font.bold

        alignment_str = "left"
        if pf and pf.alignment is not None:
            align_map = {
                0: "left",
                1: "center",
                2: "right",
                3: "both",  # justify
            }
            alignment_str = align_map.get(pf.alignment, "left")

        first_line_indent = 0
        if pf and pf.first_line_indent:
            # Convert EMU to twips
            first_line_indent = _emu_to_twips(pf.first_line_indent)

        before_spacing = 0
        if pf and pf.space_before:
            before_spacing = pf.space_before

        after_spacing = 0
        if pf and pf.space_after:
            after_spacing = pf.space_after

        struct.styles[style.name] = StyleInfo(
            style_id=style.name,
            font_name=font_name,
            font_size=font_size,
            bold=bold,
            alignment=alignment_str,
            first_line_indent=first_line_indent,
            before_spacing=before_spacing,
            after_spacing=after_spacing,
        )

    # -- 3. Page setup ----------------------------------------------------------
    try:
        sect = doc.sections[0]
        struct.page_setup.margins = MarginInfo(
            top=_emu_to_twips(_get_or_none(sect, "top_margin")),
            bottom=_emu_to_twips(_get_or_none(sect, "bottom_margin")),
            left=_emu_to_twips(_get_or_none(sect, "left_margin")),
            right=_emu_to_twips(_get_or_none(sect, "right_margin")),
        )
        struct.page_setup.page_size = PageSize(
            width=_to_int(_get_or_none(sect, "page_width")),
            height=_to_int(_get_or_none(sect, "page_height")),
        )
    except IndexError:
        pass

    # -- 4. Header / footer -----------------------------------------------------
    try:
        sect = doc.sections[0]
        header = sect.header
        if header and not header.is_linked_to_previous:
            struct.header_footer.header = "\n".join(
                p.text for p in header.paragraphs if p.text.strip()
            )

        footer = sect.footer
        if footer and not footer.is_linked_to_previous:
            struct.header_footer.footer = "\n".join(
                p.text for p in footer.paragraphs if p.text.strip()
            )
    except IndexError:
        pass

    # -- 5. Placeholders (《...》) ------------------------------------------------
    placeholder_pattern = re.compile(r"《([^》]+)》")
    texts_to_search: list[str] = []

    for para in doc.paragraphs:
        texts_to_search.append(para.text)

    texts_to_search.append(struct.header_footer.header)
    texts_to_search.append(struct.header_footer.footer)

    found: set[str] = set()
    for text in texts_to_search:
        for m in placeholder_pattern.finditer(text):
            found.add(m.group(0))
    struct.placeholders = sorted(found)

    # -- 6. Embedded images -----------------------------------------------------
    struct.images = _extract_images_from_document(doc)

    return struct


def _find_node(nodes: list[HeadingNode], level: int, text: str) -> Optional[HeadingNode]:
    """Recursively find a heading node by level and text."""
    for node in nodes:
        if node.level == level and node.text == text:
            return node
        result = _find_node(node.children, level, text)
        if result is not None:
            return result
    return None


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------


def main() -> None:
    """CLI entry point."""
    args = sys.argv[1:]
    if not args or args[0] in ("-h", "--help"):
        print(__doc__)
        sys.exit(0)

    template_path = Path(args[0])
    output_path: Optional[Path] = None

    for i, arg in enumerate(args[1:]):
        if arg == "--output" and i + 2 < len(args):
            output_path = Path(args[i + 2])

    try:
        struct = extract_template(template_path)
    except Exception as exc:
        print(f"Error extracting template: {exc}", file=sys.stderr)
        sys.exit(1)

    data = _to_dict(struct)
    json_str = json.dumps(data, ensure_ascii=False, indent=2)

    if output_path:
        output_path.write_text(json_str, encoding="utf-8")
        print(f"Written to {output_path}")
    else:
        print(json_str)


if __name__ == "__main__":
    main()
